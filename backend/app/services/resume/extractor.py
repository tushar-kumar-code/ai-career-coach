import os
import re
import pymupdf
import docx
from typing import Tuple
from fastapi import UploadFile, HTTPException

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
ALLOWED_EXTENSIONS = {".pdf", ".docx"}
UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "..", "uploads", "resumes")


class DocumentExtractor:
    """Service for validating, storing, and extracting raw text from PDF and DOCX files."""

    def __init__(self):
        os.makedirs(UPLOAD_DIR, exist_ok=True)

    async def process_uploaded_file(self, file: UploadFile, user_id: str) -> Tuple[str, str, str]:
        """
        Validates file, saves securely to disk, and extracts raw text.
        Returns: (file_path, filename, raw_text)
        """
        if not file.filename:
            raise HTTPException(status_code=400, detail="Filename is empty")

        filename = os.path.basename(file.filename)
        ext = os.path.splitext(filename)[1].lower()

        if ext not in ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file format '{ext}'. Only PDF (.pdf) and Word (.docx) files are supported."
            )

        content = await file.read()
        if len(content) == 0:
            raise HTTPException(status_code=400, detail="Uploaded file is empty (0 bytes)")

        if len(content) > MAX_FILE_SIZE:
            raise HTTPException(status_code=400, detail="File size exceeds maximum limit of 10MB")

        # Save file securely
        safe_filename = f"{user_id}_{re.sub(r'[^a-zA-Z0-9_.-]', '_', filename)}"
        file_path = os.path.join(UPLOAD_DIR, safe_filename)

        with open(file_path, "wb") as f:
            f.write(content)

        # Extract text based on file format
        raw_text = ""
        try:
            if ext == ".pdf":
                raw_text = self._extract_text_from_pdf(file_path)
            elif ext == ".docx":
                raw_text = self._extract_text_from_docx(file_path)
        except Exception as e:
            raise HTTPException(
                status_code=400,
                detail=f"Failed to read document contents. File may be corrupted or password-protected. Error: {str(e)}"
            )

        raw_text_clean = raw_text.strip()
        if len(raw_text_clean) < 20:
            raise HTTPException(
                status_code=400,
                detail="Extracted text is too short or empty. Please upload a document containing readable text."
            )

        return file_path, filename, raw_text_clean

    def _extract_text_from_pdf(self, path: str) -> str:
        text_content = []
        doc = pymupdf.open(path)
        for page in doc:
            text_content.append(page.get_text())
        doc.close()
        return "\n".join(text_content)

    def _extract_text_from_docx(self, path: str) -> str:
        doc = docx.Document(path)
        text_content = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    text_content.append(row_text)
        return "\n".join(text_content)
